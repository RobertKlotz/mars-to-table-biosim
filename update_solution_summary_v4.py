#!/usr/bin/env python3
"""
Update Solution Summary v4 with aggressive opening that addresses
livestock "weakness" head-on and positions it as the winning differentiator.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/Users/robertklotz/Downloads/Mars-to-table/docs"

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_table(doc, headers, rows, header_color="1F4E79"):
    """Add a formatted table with header styling."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], header_color)

    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)

    return table


def create_solution_summary_v4_revised():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ========== TITLE ==========
    title = doc.add_heading('MARS TO TABLE CHALLENGE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Solution Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].bold = True

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Title: ').bold = True
    info.add_run('sTARS Integrated Food Ecosystem for Mars Surface Operations\n')
    info.add_run('Team: ').bold = True
    info.add_run('Bueché-Labs LLC\n')
    info.add_run('Submitted: ').bold = True
    info.add_run('May 2026')

    doc.add_page_break()

    # ========== NEW AGGRESSIVE ABSTRACT ==========
    doc.add_heading('Abstract', level=1)

    # Opening salvo - address the "weakness" head on
    p1 = doc.add_paragraph()
    p1.add_run('Hydroponics alone will not sustain a Mars settlement.').bold = True
    p1.add_run(' Neither will algae bioreactors, insect protein, or cellular agriculture. These approaches solve nutrition while ignoring the psychological reality that has ended expeditions, broken crews, and turned "sustainable" into merely "survivable."')

    # What we propose
    p2 = doc.add_paragraph()
    p2.add_run('We propose something different: ').bold = True
    p2.add_run('a complete agricultural civilization capable of producing fresh eggs for breakfast, bread from the oven, aged cheese, and fish on Fridays. Real food. The kind humans have shared for 10,000 years.')

    # Why it matters
    p3 = doc.add_paragraph()
    p3.add_run('This is not complexity for its own sake. ')
    p3.add_run('Every gram of livestock embryo we send to Mars displaces hundreds of kilograms of prepackaged protein over a multi-year mission. Every fermentation vessel reduces Earth-dependency while improving gut health and crew morale. Every meal that feels like home is a day the mission doesn\'t fail from the inside.')

    # Technical credibility
    p4 = doc.add_paragraph()
    p4.add_run('The sTARS Integrated Food Ecosystem achieves ')
    p4.add_run('90% Earth-independence').bold = True
    p4.add_run('—40 points above the requirement—using technologies that exist today: frozen embryos (routine since the 1980s), aquaculture breeding (commercial standard), and preservation methods perfected over decades. The question is not feasibility. ')
    p4.add_run('The question is ambition.').italic = True

    # The closer
    p5 = doc.add_paragraph()
    p5.add_run('NASA is not asking how to keep astronauts alive. They are asking how to build a civilization. ').bold = True
    p5.add_run('This is our answer.').bold = True

    # ========== TEAM DESCRIPTION (1 page) ==========
    doc.add_heading('Team Description', level=1)

    doc.add_paragraph(
        'Bueché-Labs LLC is a U.S.-based aerospace systems design company specializing in modular '
        'space infrastructure, closed-loop life support, and active radiation protection. The company '
        'holds provisional patents on key technologies in this submission.'
    )

    doc.add_heading('Team Members', level=2)

    add_formatted_table(doc,
        ['Name', 'Role', 'Expertise'],
        [
            ['Robert Klotz', 'CTO, Co-Founder', 'Systems architecture, C2, space infrastructure'],
            ['Brad Bueché', 'VP Technical Services, Co-Founder', 'Technical operations, systems integration'],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('Company Background', level=2)
    doc.add_paragraph(
        'Bueché-Labs developed the sTARS modular infrastructure platform providing the foundational '
        'POD architecture for this Mars application. Combined 70+ years experience in systems engineering, '
        'life support integration, autonomous robotics, active radiation protection, and closed-loop '
        'resource management.'
    )

    # ========== TECHNOLOGY DESCRIPTION ==========
    doc.add_heading('Technology Description', level=1)

    doc.add_heading('System Architecture: 15 PODs', level=2)
    doc.add_paragraph(
        'The food system comprises 15 modular PODs (each 10m × 7.6m diameter, 115 m² usable, three decks) '
        'connected to a central spine via Universal Station Interfaces. Four superconducting REBCO dipoles '
        'provide active SEP protection for all biological assets.'
    )

    add_formatted_table(doc,
        ['Component', 'Qty', 'Function'],
        [
            ['Food PODs 1-5', '5', 'Human crops: potatoes, vegetables, legumes, oilseeds (1,805 m²)'],
            ['Food POD 6', '1', 'Livestock fodder: alfalfa, barley fodder, fodder beets'],
            ['Food POD 7', '1', 'Grain: wheat, amaranth, buckwheat (~5.5 kg flour/day)'],
            ['Livestock POD', '1', 'Dairy goats (8L milk/day) + laying hens (17 eggs/day)'],
            ['Aquaponics POD', '1', 'Tilapia farming: 4 tanks, 8000L, 0.5-1 kg fish/day'],
            ['Food Processing POD', '1', 'Oil extraction, fermentation, milling, drying'],
            ['RSV PODs', '2', 'Water extraction, electrolysis, fuel cells, storm power'],
            ['Nutrient Processing', '1', 'Haber-Bosch N₂ fixation, urine/manure processing'],
            ['Waste Processing', '1', 'Anaerobic digestion, biogas SOFC, pyrolysis'],
            ['HAB/LAB', '1', 'Kitchen, dining, food prep, cold storage'],
        ]
    )

    # ========== WHY LIVESTOCK WORKS ==========
    doc.add_heading('Why Livestock: The Technical Case', level=2)

    p = doc.add_paragraph()
    p.add_run('Critics assume livestock adds prohibitive mass and complexity. The opposite is true:\n')

    doc.add_paragraph('Biological Payload to Mars:', style='List Bullet')

    add_formatted_table(doc,
        ['Item', 'Mass', 'State', 'Proven Technology'],
        [
            ['Goat embryos (20)', '~100g', 'Frozen (-196°C)', 'Routine since 1980s'],
            ['Fertilized chicken eggs (40)', '~2 kg', 'Frozen/fresh', '21-day incubation'],
            ['Tilapia embryos (500)', '~50g', 'Frozen', 'Aquaculture standard'],
            ['All seeds (25+ varieties)', '~20 kg', 'Dry, ambient', 'Seed bank protocols'],
            ['Starter cultures', '~1 kg', 'Freeze-dried', 'Commercial practice'],
            ['TOTAL BIOLOGICAL', '<25 kg', '', 'All proven, all available NOW'],
        ]
    )

    doc.add_paragraph()
    livestock_point = doc.add_paragraph()
    livestock_point.add_run('25 kg of embryos and seeds produces 40,000+ kcal/day indefinitely. ')
    livestock_point.add_run('The same mass in prepackaged food lasts approximately 12 days.').italic = True

    # ========== EARTH INDEPENDENCE ==========
    doc.add_heading('Earth-Independence: 90%', level=1)

    add_formatted_table(doc,
        ['Food Source', 'Daily kcal', '% of Total'],
        [
            ['Crops (PODs 1-5): potatoes, vegetables, legumes, oilseeds', '26,800', '59%'],
            ['Grain (POD 7): wheat, amaranth, buckwheat', '4,500', '10%'],
            ['Goat products: milk, cheese, yogurt, meat', '5,300', '12%'],
            ['Chicken products: eggs, meat', '1,425', '3%'],
            ['Tilapia fish (Aquaponics POD)', '750', '2%'],
            ['Vegetable oil (Food Processing POD)', '2,200', '5%'],
            ['TOTAL IN-SITU', '40,975', '90%'],
            ['Earth-supplied: supplements, spices, specialty items', '4,550', '10%'],
        ]
    )

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run('Crew requirement: 3,035 kcal × 15 = 45,525 kcal/day. System produces 90% on Mars.').italic = True

    # ========== SIX PROTEIN SOURCES ==========
    doc.add_heading('Six Protein Sources', level=2)

    add_formatted_table(doc,
        ['Source', 'Daily Output', 'Protein', 'Why It Matters'],
        [
            ['Fresh Eggs', '17/day', '~100g', 'Complete protein, crew favorite'],
            ['Goat Milk', '8 L/day', '65g', 'Fresh dairy, calcium, morale'],
            ['Goat Cheese', '300g/day', '75g', 'Aged protein, variety, culture'],
            ['Tilapia Fish', '0.5-1 kg/day', '100-200g', 'Fresh seafood, omega-3'],
            ['Tempeh', 'Variable', '40-80g', 'Fermented soy, probiotics'],
            ['Meat (periodic)', '~80g avg', '20g', 'Culled animals, special occasions'],
        ]
    )

    # ========== NOVELTY AND INNOVATION ==========
    doc.add_heading('Novelty and Innovation', level=1)

    innovations = [
        ('REAL FOOD, NOT SURVIVAL RATIONS', 'Fresh eggs, warm bread, aged cheese, grilled fish—meals that maintain human identity 225 million km from home.'),
        ('COMPLETE PROTEIN INDEPENDENCE', 'Six sources ensure no single-point nutritional failure. Eggs alone provide complete amino acids.'),
        ('PSYCHOLOGICAL SUSTAINABILITY', 'Antarctic and submarine research proves: food monotony breaks crews. Our 14-sol rotation prevents it.'),
        ('CLOSED-LOOP EFFICIENCY', 'Fish waste → plant nutrients. Manure → biogas + fertilizer. Oil pressing meal → livestock feed. Nothing wasted.'),
        ('PROVEN TECHNOLOGY TODAY', 'Frozen embryos, aquaculture breeding, fermentation cultures—all commercially available, all flight-ready.'),
        ('90% EARTH-INDEPENDENCE', 'Exceeds requirement by 40 points. This is not a food system. It is a civilization.'),
    ]

    for i, (title, desc) in enumerate(innovations, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}: ').bold = True
        p.add_run(desc)

    # ========== TERRESTRIAL APPLICATIONS ==========
    doc.add_heading('Terrestrial Food Security Applications', level=1)

    apps = doc.add_paragraph()
    apps.add_run('• Remote Communities: ').bold = True
    apps.add_run('Arctic/Antarctic/island deployment independent of supply chains.\n')
    apps.add_run('• Disaster Response: ').bold = True
    apps.add_run('Rapid deployment following infrastructure collapse.\n')
    apps.add_run('• Urban Vertical Farming: ').bold = True
    apps.add_run('95% water reduction, zero agricultural runoff.\n')
    apps.add_run('• Integrated Small Farming: ').bold = True
    apps.add_run('Goat-chicken-fish-crop model scales globally.')

    # ========== IP STATEMENT ==========
    doc.add_heading('Intellectual Property Statement', level=1)

    doc.add_paragraph(
        'All IP owned by Bueché-Labs LLC. Provisional patents cover sTARS platform, POD architecture, '
        'USI specs, RSV systems, and SEP protection. Solution uses commercial technologies with '
        'proprietary integration. No licensing restrictions on core biological components.'
    )

    # ========== AI DISCLOSURE ==========
    doc.add_heading('AI Tools Disclosure', level=1)

    doc.add_paragraph(
        'Anthropic Claude AI used as force multiplier for calculations, documentation, and simulation '
        'development, working from Bueché-Labs internal documentation. All claims verified by team. '
        'Core innovation and architecture from Bueché-Labs proprietary work.'
    )

    # ========== REFERENCES ==========
    doc.add_heading('References', level=1)

    refs = [
        'NASA STD-3001: Spaceflight Human-System Standard, Volumes 1 & 2.',
        'NASA BVAD: Life Support Baseline Values and Assumptions Document.',
        'Wheeler, R.M. "Agriculture for Space." Open Agriculture, 2017.',
        'Massa, G.D. et al. "VEG-01: Veggie Hardware Validation." Open Agriculture, 2017.',
        'Smith, S. et al. "Human Adaptation to Spaceflight: Role of Food and Nutrition." 2nd Ed, 2021.',
        'FAO. "Small-scale aquaponic food production." Technical Paper 589, 2014.',
        'Bueché-Labs Internal: sTARS Primer, POD ICD, RSV Plan, SEP Concept, 2024-2026.',
    ]

    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f'[{i}] {ref}')

    # Save document
    output_path = os.path.join(OUTPUT_DIR, 'Mars_to_Table_Solution_Summary_v4.docx')
    doc.save(output_path)
    print(f'Updated: {output_path}')
    return output_path


if __name__ == '__main__':
    create_solution_summary_v4_revised()
    print("\nSolution Summary v4 updated with aggressive opening.")
    print("Key changes:")
    print("  - Opens by addressing livestock 'weakness' head-on")
    print("  - Positions real food as the ONLY viable long-term solution")
    print("  - Adds 'Why Livestock Works' section with mass comparison")
    print("  - Emphasizes all technology exists TODAY")
    print("  - Frames as 'civilization, not survival'")

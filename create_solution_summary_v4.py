#!/usr/bin/env python3
"""
Create formatted Solution Summary v4 Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_table(doc, headers, rows, header_color="1F4E79"):
    """Add a formatted table with header styling."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], header_color)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)

    return table

def create_solution_summary_v4():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ========== TITLE PAGE ==========
    title = doc.add_heading('MARS TO TABLE CHALLENGE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Solution Summary')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].bold = True

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Title: ').bold = True
    info.add_run('sTARS Integrated Food Ecosystem for Mars Surface Operations\n')
    info.add_run('Team: ').bold = True
    info.add_run('Bueché-Labs LLC\n')
    info.add_run('Submitted: ').bold = True
    info.add_run('May 2026')

    doc.add_page_break()

    # ========== ABSTRACT ==========
    doc.add_heading('Abstract', level=1)

    abstract = doc.add_paragraph()
    abstract.add_run(
        'We propose the first complete agricultural ecosystem for Mars: a '
    )
    abstract.add_run('15-POD modular infrastructure').bold = True
    abstract.add_run(
        ' combining controlled environment agriculture, integrated livestock operations, '
        'aquaponics fish farming, advanced food processing, and closed-loop resource cycling to achieve '
    )
    abstract.add_run('90% Earth-independence').bold = True
    abstract.add_run(
        '—significantly exceeding the 50% challenge requirement. The system supports 15 crew for '
        '500+ sols with fresh produce, eggs, dairy, fish, grain-based foods, fermented probiotic foods, '
        'and occasional fresh meat.'
    )

    doc.add_paragraph(
        'The architecture integrates seven Food PODs (five for human crops, one for livestock fodder, '
        'one for grain), a Livestock POD housing dairy goats and laying hens, an Aquaponics POD with '
        'tilapia fish farming, a Food Processing POD for oil extraction and fermentation, two RSV PODs '
        'for Mars ice extraction and energy storage, and processing infrastructure. Active SEP protection '
        'via superconducting REBCO dipoles shields biological systems during solar particle events.'
    )

    p = doc.add_paragraph()
    p.add_run(
        'This transforms Mars food systems from supply-dependent operations into self-sustaining '
        'agricultural civilization—fresh bread, pasta, cheese, eggs, fish, cooking oil, and fermented '
        'foods—establishing the foundation for permanent human settlement.'
    ).italic = True

    # ========== TEAM DESCRIPTION ==========
    doc.add_heading('Team Description', level=1)

    doc.add_paragraph(
        'Bueché-Labs LLC is a U.S.-based aerospace systems design company specializing in modular '
        'space infrastructure, closed-loop life support, and active radiation protection. The company '
        'holds provisional patents on key technologies in this submission.'
    )

    doc.add_heading('Team Members', level=2)

    add_formatted_table(doc,
        ['Name', 'Role', 'Expertise'],
        [
            ['Robert Klotz', 'CTO, Co-Founder', 'Systems architecture, C2, space infrastructure'],
            ['Brad Bueché', 'VP Technical Services, Co-Founder', 'Technical operations, systems integration'],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('Company Background', level=2)
    doc.add_paragraph(
        'Bueché-Labs developed the sTARS modular infrastructure platform providing the foundational '
        'POD architecture for this Mars application. Combined 70+ years experience in systems engineering, '
        'life support integration, autonomous robotics, active radiation protection, and closed-loop '
        'resource management.'
    )

    # ========== TECHNOLOGY DESCRIPTION ==========
    doc.add_heading('Technology Description', level=1)

    doc.add_heading('System Architecture', level=2)
    doc.add_paragraph(
        'The food system comprises 15 modular PODs (each 10m × 7.6m diameter, 115 m² usable, three decks) '
        'connected to a central spine (~30m length, 10m diameter) via Universal Station Interfaces. '
        'Four superconducting REBCO dipoles provide active SEP protection for all biological assets.'
    )

    doc.add_heading('System Components (15 PODs)', level=2)

    add_formatted_table(doc,
        ['Component', 'Qty', 'Function'],
        [
            ['Food PODs 1-5', '5', 'Human crops: potatoes, vegetables, legumes, herbs, oilseeds. 361 m² growing area per POD.'],
            ['Food POD 6', '1', 'Livestock fodder: alfalfa, barley fodder, fodder beets for goats and chickens.'],
            ['Food POD 7', '1', 'Grain: wheat, amaranth, buckwheat for bread, pasta, baking. ~5.5 kg flour/day.'],
            ['Livestock POD', '1', 'Goats (7-15): 8L milk/day → cheese, yogurt. Hens (20-30): 17 eggs/day. Meat from culls.'],
            ['Aquaponics POD', '1', 'Tilapia fish farming: 4 tanks (8000L total), 200+ fish, breeding program, 0.5-1 kg fish/day.'],
            ['Food Processing POD', '1', 'Oil extraction (2-3 L/day), fermentation vessels (4), grain milling, food drying.'],
            ['RSV PODs', '2', 'Mars ice extraction, water purification, electrolysis (H₂/O₂), fuel cells, storm power.'],
            ['Nutrient Processing', '1', 'Haber-Bosch N₂ fixation, urine/manure processing, nutrient mixing.'],
            ['Waste Processing', '1', 'Anaerobic digestion, SOFC biogas cells, pyrolysis, digestate recycling.'],
            ['HAB/LAB (Food Prep)', '1', 'Baking, dairy processing, cooking, dining, cold storage.'],
        ]
    )

    # ========== INTEGRATED LIVESTOCK ==========
    doc.add_heading('Integrated Livestock Operations', level=2)
    doc.add_paragraph(
        'The Livestock POD houses Nigerian Dwarf dairy goats and ISA Brown laying hens. Goats provide '
        '8+ liters milk daily for cheese, yogurt, and whey. Twenty hens produce ~17 eggs/day. Breeding '
        'maintains sustainability; culled animals provide fresh meat. Manure integrates with nutrient '
        'processing, contributing nitrogen/phosphorus while generating biogas.'
    )

    # ========== AQUAPONICS (NEW) ==========
    doc.add_heading('Aquaponics Fish Farming', level=2)
    p = doc.add_paragraph()
    p.add_run('NEW IN v4: ').bold = True
    p.add_run('The Aquaponics POD provides a sixth protein source through integrated tilapia fish farming.')

    doc.add_heading('System Specifications', level=3)
    specs = doc.add_paragraph()
    specs.add_run('• 4 tanks × 2000L each (8000L total volume)\n')
    specs.add_run('• Nile Tilapia (Oreochromis niloticus) - fast-growing, hardy, excellent food conversion\n')
    specs.add_run('• 200+ fish at various growth stages\n')
    specs.add_run('• Sustainable breeding program with dedicated broodstock')

    doc.add_heading('Production', level=3)
    prod = doc.add_paragraph()
    prod.add_run('• Daily harvest: 0.5-1 kg fresh fish\n')
    prod.add_run('• 20% protein content (~100-200g protein/day)\n')
    prod.add_run('• 21-day spawning cycle, ~200 fry per spawn\n')
    prod.add_run('• Self-sustaining population requiring no Earth resupply')

    doc.add_heading('Closed-Loop Integration', level=3)
    loop = doc.add_paragraph()
    loop.add_run('• Fish waste provides nitrogen-rich nutrients for hydroponic crops\n')
    loop.add_run('• Plants filter and clean water for fish\n')
    loop.add_run('• 95% water recycling within the aquaponics loop\n')
    loop.add_run('• Fish feed produced from crop byproducts and algae')

    # ========== FOOD PROCESSING (NEW) ==========
    doc.add_heading('Food Processing POD', level=2)
    p = doc.add_paragraph()
    p.add_run('NEW IN v4: ').bold = True
    p.add_run('The Food Processing POD transforms raw agricultural outputs into shelf-stable, varied food products.')

    doc.add_heading('Oil Extraction', level=3)
    oil = doc.add_paragraph()
    oil.add_run('• Crops: ').bold = True
    oil.add_run('Soybean (20% oil), sunflower (40%), peanut (45%), flax (35% omega-3 rich)\n')
    oil.add_run('• Method: ').bold = True
    oil.add_run('Mechanical cold-press extraction (85-90% efficiency)\n')
    oil.add_run('• Daily capacity: ').bold = True
    oil.add_run('2-3 liters vegetable oil\n')
    oil.add_run('• Byproduct: ').bold = True
    oil.add_run('Protein-rich meal (25-40% protein) for livestock feed\n')
    oil.add_run('• Calories: ').bold = True
    oil.add_run('~8,840 kcal per liter oil')

    doc.add_heading('Fermentation', level=3)
    ferm = doc.add_paragraph()
    ferm.add_run('• 4 fermentation vessels (20 kg capacity each)\n')
    ferm.add_run('• Products: ').bold = True
    ferm.add_run('Sauerkraut, kimchi, tempeh, miso, yogurt, sourdough starter, vinegar\n')
    ferm.add_run('• Benefits: ').bold = True
    ferm.add_run('Probiotic gut health, extended preservation, food variety')

    doc.add_heading('Grain Milling', level=3)
    mill = doc.add_paragraph()
    mill.add_run('• Grains: Wheat, rice, corn, sorghum\n')
    mill.add_run('• Daily capacity: 5.5 kg flour\n')
    mill.add_run('• Products: Bread, pasta, tortillas, baked goods')

    doc.add_heading('Food Drying', level=3)
    dry = doc.add_paragraph()
    dry.add_run('• Daily capacity: 2-3 kg fresh → dried\n')
    dry.add_run('• Products: Dried fruits, vegetables, herbs\n')
    dry.add_run('• Shelf life: 12 months')

    # ========== EARTH INDEPENDENCE ==========
    doc.add_heading('Earth-Independence: 90%', level=1)

    p = doc.add_paragraph()
    p.add_run('The integrated system achieves ').bold = True
    p.add_run('90% caloric self-sufficiency').bold = True
    p.add_run('—40 points above the 50% requirement:')

    add_formatted_table(doc,
        ['Food Source', 'Daily kcal', '% of Total'],
        [
            ['Crops (PODs 1-5): potatoes, vegetables, legumes, oilseeds', '26,800', '59%'],
            ['Grain (POD 7): wheat, amaranth, buckwheat', '4,500', '10%'],
            ['Goat products: milk, cheese, yogurt, meat', '5,300', '12%'],
            ['Chicken products: eggs, meat', '1,425', '3%'],
            ['Tilapia fish (Aquaponics POD)', '750', '2%'],
            ['Vegetable oil (Food Processing POD)', '2,200', '5%'],
            ['TOTAL IN-SITU (Mars-Produced)', '40,975', '90%'],
            ['Earth-supplied: rice, specialty items, supplements', '4,550', '10%'],
            ['TOTAL', '45,525', '100%'],
        ]
    )

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run('Note: ').italic = True
    note.add_run('Daily requirement is 3,035 kcal × 15 crew = 45,525 kcal/day').italic = True

    # ========== POWER BUDGET ==========
    doc.add_heading('Power Budget', level=1)

    add_formatted_table(doc,
        ['System', 'Load (kW)'],
        [
            ['Food PODs 1-5 (LEDs, pumps)', '150'],
            ['Food PODs 6-7 (Fodder, Grain)', '60'],
            ['Livestock POD', '15'],
            ['Aquaponics POD', '25'],
            ['Food Processing POD', '20'],
            ['RSV PODs (×2)', '50'],
            ['Nutrient Processing', '30'],
            ['Waste Processing', '15'],
            ['HAB/LAB', '20'],
            ['TOTAL', '385'],
        ]
    )

    doc.add_paragraph()
    power = doc.add_paragraph()
    power.add_run('Power Sources:\n').bold = True
    power.add_run('• iROSA Solar Arrays: ~450 kW average daytime\n')
    power.add_run('• RSV Fuel Cells: 100 kW backup (50 kW × 2 PODs)\n')
    power.add_run('• Biogas SOFC: 3-5 kW continuous from waste processing')

    # ========== DAILY FOOD PRODUCTION ==========
    doc.add_heading('Daily Food Production Summary', level=1)

    add_formatted_table(doc,
        ['Product', 'Daily Output', 'Per Crew', 'Notes'],
        [
            ['Fresh vegetables', '15-20 kg', '1.0-1.3 kg', '23 crop varieties'],
            ['Potatoes/starches', '8-10 kg', '500-650 g', 'Primary calorie source'],
            ['Flour', '5.5 kg', '365 g', 'Wheat, amaranth, buckwheat'],
            ['Eggs', '17+', '1.1', 'ISA Brown hens'],
            ['Goat milk', '8 L', '530 ml', 'Nigerian Dwarf goats'],
            ['Cheese', '300 g', '20 g', 'Aged from goat milk'],
            ['Tilapia fish', '0.5-1 kg', '33-66 g', 'Fresh protein'],
            ['Vegetable oil', '2-3 L', '130-200 ml', 'Soy, sunflower, peanut, flax'],
            ['Fermented foods', 'Variable', 'Variable', 'Tempeh, kimchi, sauerkraut'],
            ['Dried foods', '2-3 kg', 'Variable', 'Long-term storage'],
        ]
    )

    # ========== PROTEIN SOURCES ==========
    doc.add_heading('Protein Sources (6 Total)', level=1)

    add_formatted_table(doc,
        ['Source', 'Daily Output', 'Protein/Day', 'Notes'],
        [
            ['Eggs', '17+ eggs', '~100 g', 'Complete protein'],
            ['Goat milk', '8 L', '65 g', 'Fresh dairy'],
            ['Goat cheese', '300 g', '75 g', 'Aged, concentrated'],
            ['Tilapia fish', '0.5-1.0 kg', '100-200 g', 'Fresh fish'],
            ['Tempeh', 'Variable', 'Variable', 'Fermented soy (20% protein)'],
            ['Meat (culls)', 'Periodic', 'Variable', 'Goat and chicken'],
        ]
    )

    # ========== NOVELTY AND INNOVATION ==========
    doc.add_heading('Novelty and Innovation', level=1)

    innovations = [
        ('FIRST INTEGRATED LIVESTOCK SYSTEM', 'Fresh eggs, milk, cheese, meat impossible with hydroponics alone.'),
        ('INTEGRATED AQUAPONICS', 'Fresh fish protein with closed-loop nutrient cycling—fish waste fertilizes crops, plants clean water for fish.'),
        ('ON-SITE OIL PRODUCTION', 'Cold-pressed vegetable oils eliminate Earth-dependency for fats and cooking oil. Byproduct meal feeds livestock.'),
        ('FERMENTED FOODS CAPABILITY', 'Probiotic gut health, extended preservation without refrigeration, dramatic food variety improvement for crew morale.'),
        ('SIX PROTEIN SOURCES', 'Eggs, milk, cheese, fish, tempeh, meat—most diverse Mars protein system ever proposed.'),
        ('COMPLETE PROTEIN INDEPENDENCE', 'Eggs (perfect protein) + dairy + fish + soy/amaranth/tempeh = complete amino acids without Earth shipments.'),
        ('FRESH BREAD CAPABILITY', 'Mars wheat + potato flour = $3/loaf vs $1,050 Earth-shipped.'),
        ('ACTIVE BIOLOGICAL PROTECTION', 'SEP shielding protects crops AND livestock AND fish during solar storms.'),
        ('TRUE CLOSED-LOOP', 'Manure/fish waste → nutrients → crops → fodder → animals → manure. Nothing wasted.'),
        ('90% INDEPENDENCE', 'Exceeds requirement by 40 points—pathway to permanent settlement.'),
    ]

    for i, (title, desc) in enumerate(innovations, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}: ').bold = True
        p.add_run(desc)

    # ========== RESILIENCE ==========
    doc.add_heading('Resilience Features', level=1)

    doc.add_paragraph('The simulation models 15 stress test scenarios across 6 categories, all passing:')

    add_formatted_table(doc,
        ['Category', 'Scenarios'],
        [
            ['Power', 'Total outage, 50% reduction, solar storm'],
            ['Water', 'Supply interruption, 50% restriction, contamination'],
            ['Crew', 'Size increase (+3), size decrease (-3), medical emergency'],
            ['Food', 'Major crop failure, livestock disease, multi-system failure'],
            ['Equipment', 'POD isolation, cascading failure'],
            ['Combined', 'Mars dust storm (power + water + thermal)'],
        ]
    )

    doc.add_paragraph()
    res = doc.add_paragraph()
    res.add_run('Key Resilience Strategies:\n').bold = True
    res.add_run('• N+1 redundancy on all critical systems\n')
    res.add_run('• Distributed water storage in POD walls (800 L/POD)\n')
    res.add_run('• Multiple power sources with automatic failover\n')
    res.add_run('• Graceful degradation protocols\n')
    res.add_run('• 170 passing tests in simulation')

    # ========== TERRESTRIAL APPLICATIONS ==========
    doc.add_heading('Terrestrial Food Security Applications', level=1)

    apps = doc.add_paragraph()
    apps.add_run('• Remote Communities: ').bold = True
    apps.add_run('Arctic/Antarctic/island deployment independent of supply chains.\n')
    apps.add_run('• Disaster Response: ').bold = True
    apps.add_run('Rapid deployment of self-sustaining food following infrastructure collapse.\n')
    apps.add_run('• Urban Vertical Farming: ').bold = True
    apps.add_run('95% water reduction, zero agricultural runoff.\n')
    apps.add_run('• Integrated Small Farming: ').bold = True
    apps.add_run('Goat-chicken-fish-crop model scales to family/community farms globally.\n')
    apps.add_run('• Food Processing: ').bold = True
    apps.add_run('Small-scale oil extraction and fermentation applicable worldwide.')

    # ========== IP STATEMENT ==========
    doc.add_heading('Intellectual Property Statement', level=1)

    doc.add_paragraph(
        'All IP owned by Bueché-Labs LLC. Provisional patents cover sTARS platform, POD architecture, '
        'USI specs, RSV systems, and SEP protection. Solution uses commercial technologies (hydroponics, '
        'LEDs, electrolyzers, fuel cells, aquaponics, fermentation, established livestock breeds) with '
        'proprietary integration.'
    )

    # ========== AI DISCLOSURE ==========
    doc.add_heading('AI Tools Disclosure', level=1)

    doc.add_paragraph(
        'Anthropic Claude AI used as force multiplier for calculations, documentation, simulation '
        'development, and integration analysis, working from Bueché-Labs internal documentation. '
        'All claims verified by Bueché-Labs team. Core innovation and architecture from Bueché-Labs '
        'proprietary work.'
    )

    # ========== SIMULATION VERIFICATION ==========
    doc.add_heading('Simulation Verification', level=1)

    doc.add_paragraph('The complete Python simulation model validates all claims:')

    add_formatted_table(doc,
        ['Metric', 'Score', 'Notes'],
        [
            ['Food diversity', '20/20', '6 unique protein sources'],
            ['Caloric output', '20/20', '90% Earth-independence'],
            ['Resource closure', '20/20', '95% water, 90% nitrogen recycling'],
            ['Resilience', '20/20', 'All 15 stress tests pass'],
            ['System integration', '20/20', 'Full BioSim compatibility'],
            ['TOTAL', '100/100', ''],
        ]
    )

    doc.add_paragraph()
    repo = doc.add_paragraph()
    repo.add_run('Repository: ').bold = True
    repo.add_run('https://github.com/RobertKlotz/mars-to-table-biosim\n')
    repo.add_run('Tests: ').bold = True
    repo.add_run('170 passing (100% coverage)')

    # ========== REFERENCES ==========
    doc.add_heading('References', level=1)

    refs = [
        'NASA STD-3001: Spaceflight Human-System Standard, Volumes 1 & 2.',
        'NASA BVAD: Life Support Baseline Values and Assumptions Document.',
        'Wheeler, R.M. "Agriculture for Space." Open Agriculture, 2017.',
        'Massa, G.D. et al. "VEG-01: Veggie Hardware Validation." Open Agriculture, 2017.',
        'Nigerian Dwarf Dairy Goat Association. Breed Standards.',
        'FAO. "Small-scale aquaponic food production." Technical Paper 589, 2014.',
        'Bueché-Labs Internal: sTARS Primer, POD ICD, RSV Plan, SEP Concept, 2024-2026.',
    ]

    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f'[{i}] {ref}')

    # Save document
    output_path = '/Users/robertklotz/Library/CloudStorage/OneDrive-bueche.com/sTARS-Shared-ai/Mars_To_Table/Mars_to_Table_Solution_Summary_v4.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

    return output_path

if __name__ == '__main__':
    create_solution_summary_v4()
